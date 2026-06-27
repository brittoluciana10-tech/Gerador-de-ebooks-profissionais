import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors
from datetime import datetime
from io import BytesIO
import requests

st.set_page_config(
    page_title="Ebook Creator Pro Premium",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# CSS DESIGN PROFISSIONAL
st.markdown("""
    <style>
    * { margin: 0; padding: 0; box-sizing: border-box; }
    
    [data-testid="stAppViewContainer"] {
        background: linear-gradient(135deg, #0a0e27 0%, #1a1f3a 50%, #0f172a 100%);
        color: #f1f5f9;
    }
    
    [data-testid="stHeader"] { background: transparent !important; }
    
    .stButton>button {
        background: linear-gradient(135deg, #6366f1 0%, #4f46e5 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 10px !important;
        padding: 14px 28px !important;
        font-weight: 700 !important;
        font-size: 16px !important;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
        box-shadow: 0 8px 24px rgba(99, 102, 241, 0.35) !important;
        letter-spacing: 0.5px !important;
    }
    
    .stButton>button:hover {
        transform: translateY(-4px) !important;
        box-shadow: 0 12px 32px rgba(99, 102, 241, 0.5) !important;
        background: linear-gradient(135deg, #4f46e5 0%, #4338ca 100%) !important;
    }
    
    .header-premium {
        background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 50%, #a855f7 100%);
        border-radius: 20px;
        padding: 48px 32px;
        margin: 24px 0;
        text-align: center;
        box-shadow: 0 20px 60px rgba(99, 102, 241, 0.3);
        border: 1px solid rgba(255, 255, 255, 0.1);
    }
    
    .header-premium h1 {
        font-size: 52px !important;
        font-weight: 800 !important;
        margin: 0 !important;
        color: white !important;
        text-shadow: 0 2px 10px rgba(0, 0, 0, 0.3);
        letter-spacing: -1px;
    }
    
    .header-premium p {
        font-size: 20px !important;
        color: rgba(255, 255, 255, 0.95) !important;
        margin: 12px 0 0 0 !important;
        font-weight: 500;
    }
    
    .section-header {
        display: flex;
        align-items: center;
        gap: 12px;
        margin: 32px 0 20px 0;
        padding-bottom: 16px;
        border-bottom: 2px solid rgba(99, 102, 241, 0.3);
    }
    
    .section-header h2 {
        color: #6366f1 !important;
        font-size: 22px !important;
        font-weight: 700 !important;
        margin: 0 !important;
    }
    
    .divider-premium {
        height: 2px;
        background: linear-gradient(90deg, transparent 0%, #6366f1 50%, transparent 100%);
        margin: 40px 0;
    }
    
    .info-box {
        background: linear-gradient(135deg, rgba(99, 102, 241, 0.1) 0%, rgba(139, 92, 246, 0.05) 100%);
        border-left: 4px solid #6366f1;
        border-radius: 8px;
        padding: 16px 20px;
        margin: 12px 0;
    }
    
    .success-box {
        background: linear-gradient(135deg, rgba(16, 185, 129, 0.1) 0%, rgba(5, 150, 105, 0.05) 100%);
        border-left: 4px solid #10b981;
        border-radius: 8px;
        padding: 16px 20px;
        margin: 12px 0;
    }
    
    footer { display: none !important; }
    
    .footer-custom {
        text-align: center;
        padding: 32px 24px;
        margin-top: 48px;
        background: linear-gradient(135deg, rgba(99, 102, 241, 0.05) 0%, rgba(139, 92, 246, 0.05) 100%);
        border-top: 1px solid rgba(99, 102, 241, 0.2);
        border-radius: 16px;
    }
    
    .footer-custom p {
        color: #cbd5e1 !important;
        font-size: 14px !important;
        margin: 4px 0 !important;
    }
    
    .footer-custom strong {
        color: #f1f5f9 !important;
    }
    </style>
""", unsafe_allow_html=True)

# FUNCAO: Gerar PDF
def gerar_pdf_profissional(tema, audience, idioma, regiao, conteudo, sugestoes_imagens):
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=0.8*inch,
        bottomMargin=0.8*inch,
        leftMargin=0.8*inch,
        rightMargin=0.8*inch
    )
    
    story = []
    styles = getSampleStyleSheet()
    
    titulo_capa = ParagraphStyle(
        'TituloCapa',
        parent=styles['Heading1'],
        fontSize=48,
        textColor=colors.HexColor('#6366f1'),
        spaceAfter=20,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    subtitulo_capa = ParagraphStyle(
        'SubtituloCapa',
        parent=styles['Normal'],
        fontSize=24,
        textColor=colors.HexColor('#8b5cf6'),
        spaceAfter=40,
        alignment=TA_CENTER,
        fontName='Helvetica'
    )
    
    info_capa = ParagraphStyle(
        'InfoCapa',
        parent=styles['Normal'],
        fontSize=12,
        textColor=colors.HexColor('#64748b'),
        spaceAfter=8,
        alignment=TA_CENTER
    )
    
    heading1 = ParagraphStyle(
        'CustomHeading1',
        parent=styles['Heading1'],
        fontSize=28,
        textColor=colors.HexColor('#0f172a'),
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    heading2 = ParagraphStyle(
        'CustomHeading2',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#1e293b'),
        spaceAfter=10,
        spaceBefore=8,
        fontName='Helvetica-Bold'
    )
    
    body = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=10,
        alignment=TA_JUSTIFY,
        spaceAfter=10,
        leading=15,
        textColor=colors.HexColor('#334155')
    )
    
    # CAPA
    story.append(Spacer(1, 1.5*inch))
    story.append(Paragraph("📚", titulo_capa))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph(tema, titulo_capa))
    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph(f"Um guia essencial para {audience}", subtitulo_capa))
    story.append(Spacer(1, 1.5*inch))
    story.append(Paragraph("Por Luciana Britto | L&B Marketing", info_capa))
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph("Estratégias de Valor", info_capa))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph(f"Idioma: {idioma} | Região: {regiao}", info_capa))
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph(f"© 2026 — Todos os direitos reservados", info_capa))
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph(datetime.now().strftime("Gerado em %d de %B de %Y"), info_capa))
    
    # ÍNDICE
    story.append(PageBreak())
    story.append(Paragraph("ÍNDICE", heading1))
    story.append(Spacer(1, 0.2*inch))
    
    for idx, item in enumerate(["Introdução", "Conteúdo Principal", "Conclusão", "Sugestões de Imagens"], 1):
        story.append(Paragraph(f"{idx}. {item}", body))
    
    # CONTEÚDO
    story.append(PageBreak())
    story.append(Paragraph("CONTEÚDO PRINCIPAL", heading1))
    story.append(Spacer(1, 0.2*inch))
    
    linhas = conteudo.split('\n')
    current_chapter = 0
    
    for linha in linhas:
        linha = linha.strip()
        if not linha:
            story.append(Spacer(1, 0.08*inch))
        elif linha.startswith('###'):
            if current_chapter > 0:
                story.append(PageBreak())
            current_chapter += 1
            titulo = linha.replace('###', '').replace('**', '').strip()
            story.append(Paragraph(titulo, heading1))
            story.append(Spacer(1, 0.2*inch))
        elif linha.startswith('##'):
            titulo = linha.replace('##', '').replace('**', '').strip()
            story.append(Paragraph(titulo, heading2))
            story.append(Spacer(1, 0.1*inch))
        else:
            clean_text = linha.replace('**', '').replace('*', '').replace('`', '').strip()
            if clean_text and len(clean_text) > 2:
                try:
                    story.append(Paragraph(clean_text, body))
                except:
                    pass
    
    # IMAGENS
    story.append(PageBreak())
    story.append(Paragraph("SUGESTÕES DE IMAGENS PARA CANVA", heading1))
    story.append(Spacer(1, 0.2*inch))
    
    linhas_imagens = sugestoes_imagens.split('\n')
    for linha in linhas_imagens:
        linha = linha.strip()
        if linha and len(linha) > 2:
            clean = linha.replace('**', '').replace('*', '')
            try:
                story.append(Paragraph(clean, body))
            except:
                pass
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

# FUNCAO: Download
def fazer_download_ebook(tema, audience, idioma, regiao, conteudo, sugestoes_imagens, formato_txt, formato_word, formato_pdf):
    col_d1, col_d2, col_d3 = st.columns(3)
    
    if formato_txt:
        with col_d1:
            txt = f"{tema}\n\nPor Luciana Britto | L&B Marketing\n{idioma} | {regiao}\n\n{conteudo}\n\n{sugestoes_imagens}"
            st.download_button("📄 TXT", txt, f"{tema}.txt", "text/plain", use_container_width=True)
    
    if formato_word:
        with col_d2:
            doc = Document()
            title = doc.add_heading(tema, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(conteudo)
            doc.add_page_break()
            doc.add_heading("Imagens Sugeridas", level=1)
            doc.add_paragraph(sugestoes_imagens)
            
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            st.download_button("📋 Word", doc_bytes.getvalue(), f"{tema}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    
    if formato_pdf:
        with col_d3:
            pdf_data = gerar_pdf_profissional(tema, audience, idioma, regiao, conteudo, sugestoes_imagens)
            st.download_button("🎨 PDF", pdf_data, f"{tema}.pdf", "application/pdf", use_container_width=True)

# INICIALIZAR STATE
if "ebooks" not in st.session_state:
    st.session_state.ebooks = []

if "page" not in st.session_state:
    st.session_state.page = "onboarding"

if "selected_type" not in st.session_state:
    st.session_state.selected_type = "Documento"

# TEMAS
TEMAS = {
    "🤰 Maternidade Real": "Maternidade real, mães, experiências honestas",
    "👶 Sono do Bebé": "Sono infantil, técnicas de dormir, rotina",
    "💪 Pós-parto": "Recuperação pós-parto, saúde, bem-estar",
    "📚 Parenting": "Educação infantil, desenvolvimento, parentalidade",
    "💰 Renda Variável": "Renda extra, freelance, negócios digitais",
    "🧘 Bem-estar": "Mindfulness, meditação, saúde mental",
    "💼 Autónomo": "Ser autónomo, gestão, financeiro",
    "⚖️ Peso Saudável": "Perda de peso, fitness, saúde, nutrição",
    "🎨 Criatividade": "Criatividade, inovação, arte, inspiração",
    "💳 Finanças": "Educação financeira, investimentos, economia",
    "❤️ Relacionamentos": "Relacionamentos, amor, comunicação",
    "🚀 Carreira": "Carreira, emprego, desenvolvimento profissional",
}

IDIOMAS = {
    "🇵🇹 Português": "Português",
    "🇬🇧 Inglês": "English",
    "🇪🇸 Espanhol": "Español",
    "🇫🇷 Francês": "Français",
    "🇮🇹 Italiano": "Italiano",
}

ESTILOS_ARTE = {
    "📷 Foto": "Fotografia realista, alta qualidade",
    "🎨 Ilustração": "Ilustração artística e colorida",
    "🔷 Abstrato": "Arte abstrata e moderna",
    "🌐 Digital": "Arte digital e futurista",
    "🌊 Aquarela": "Estilo aquarela e suave",
    "⬜ Minimalista": "Design minimalista e limpo",
}

TIPOS_DOCUMENTO = {
    "Apresentação": "📊",
    "Página da Web": "🌐",
    "Documento": "📄",
    "Social": "📱",
    "Gráfico": "📈",
}

ESTILOS = {
    "Clássico": "Design clássico e profissional",
    "Moderno": "Design moderno e inovador",
    "Padrão": "Design padrão e balanceado",
    "Criativo": "Design criativo e único",
}

TEMPLATES = {
    "📖 Guia Completo": "Um guia detalhado com introdução, capítulos e conclusão",
    "📋 Check-list": "Lista de verificação, dicas e boas práticas",
    "💡 Dicas Rápidas": "10-20 dicas rápidas e práticas",
    "📊 Análise Dados": "Análise de dados, estatísticas e gráficos",
    "🎯 Plano de Ação": "Passo a passo com plano de implementação",
    "❓ FAQ": "Perguntas frequentes com respostas detalhadas",
}

# PAGE: ONBOARDING
if st.session_state.page == "onboarding":
    st.markdown("""
        <style>
        .stApp {
            background: linear-gradient(135deg, #e0f2fe 0%, #dbeafe 50%, #f0f9ff 100%) !important;
        }
        [data-testid="stAppViewContainer"] {
            background: linear-gradient(135deg, #e0f2fe 0%, #dbeafe 50%, #f0f9ff 100%) !important;
        }
        </style>
    """, unsafe_allow_html=True)
    
    st.markdown("""
        <div style="text-align: center; margin: 60px 0 40px 0;">
            <h1 style="font-size: 56px; font-weight: 800; color: #0c2340; margin: 0;">Criar com IA</h1>
            <p style="font-size: 24px; color: #1e40af; margin: 12px 0 0 0; font-weight: 500;">Como você quer começar?</p>
        </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4, gap="medium")
    
    with col1:
        st.markdown("""
            <div style="
                background: white;
                border-radius: 16px;
                overflow: hidden;
                box-shadow: 0 10px 30px rgba(0, 0, 0, 0.1);
                height: 100%;
            ">
                <div style="
                    width: 100%;
                    height: 180px;
                    background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 50%, #a855f7 100%);
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    font-size: 72px;
                ">📚</div>
                <div style="padding: 24px;">
                    <h3 style="font-size: 20px; font-weight: 700; color: #0c2340; margin: 0 0 8px 0;">Gerar</h3>
                    <p style="font-size: 14px; color: #64748b; margin: 0; line-height: 1.5;">Criar a partir de um prompt em poucos segundos</p>
                </div>
            </div>
        """, unsafe_allow_html=True)
        if st.button("Começar", key="btn_gerar", use_container_width=True):
            st.session_state.page = "criar"
            st.rerun()
    
    with col2:
        st.markdown("""
            <div style="
                background: white;
                border-radius: 16px;
                overflow: hidden;
                box-shadow: 0 10px 30px rgba(0, 0, 0, 0.1);
                height: 100%;
            ">
                <div style="
                    width: 100%;
                    height: 180px;
                    background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 50%, #a855f7 100%);
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    font-size: 72px;
                ">✏️</div>
                <div style="padding: 24px;">
                    <h3 style="font-size: 20px; font-weight: 700; color: #0c2340; margin: 0 0 8px 0;">Colar texto</h3>
                    <p style="font-size: 14px; color: #64748b; margin: 0; line-height: 1.5;">Criar a partir de anotações ou conteúdo existente</p>
                </div>
            </div>
        """, unsafe_allow_html=True)
        if st.button("Colar", key="btn_colar", use_container_width=True):
            st.session_state.page = "colar"
            st.rerun()
    
    with col3:
        st.markdown("""
            <div style="
                background: white;
                border-radius: 16px;
                overflow: hidden;
                box-shadow: 0 10px 30px rgba(0, 0, 0, 0.1);
                height: 100%;
            ">
                <div style="
                    width: 100%;
                    height: 180px;
                    background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 50%, #a855f7 100%);
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    font-size: 72px;
                ">📋</div>
                <div style="padding: 24px;">
                    <h3 style="font-size: 20px; font-weight: 700; color: #0c2340; margin: 0 0 8px 0;">Usar modelo</h3>
                    <p style="font-size: 14px; color: #64748b; margin: 0; line-height: 1.5;">Crie usando a estrutura de um modelo</p>
                </div>
            </div>
        """, unsafe_allow_html=True)
        if st.button("Modelo", key="btn_modelo", use_container_width=True):
            st.session_state.page = "modelo"
            st.rerun()
    
    with col4:
        st.markdown("""
            <div style="
                background: white;
                border-radius: 16px;
                overflow: hidden;
                box-shadow: 0 10px 30px rgba(0, 0, 0, 0.1);
                height: 100%;
            ">
                <div style="
                    width: 100%;
                    height: 180px;
                    background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 50%, #a855f7 100%);
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    font-size: 72px;
                ">📂</div>
                <div style="padding: 24px;">
                    <h3 style="font-size: 20px; font-weight: 700; color: #0c2340; margin: 0 0 8px 0;">Importar</h3>
                    <p style="font-size: 14px; color: #64748b; margin: 0; line-height: 1.5;">Aprimorar documentos ou páginas existentes</p>
                </div>
            </div>
        """, unsafe_allow_html=True)
        if st.button("Importar", key="btn_importar", use_container_width=True):
            st.session_state.page = "importar"
            st.rerun()
    
    st.markdown("""
        <div style="margin-top: 80px; text-align: center;">
            <h2 style="font-size: 28px; font-weight: 700; color: #0c2340; margin: 0 0 24px 0;">Seus prompts recentes</h2>
        </div>
    """, unsafe_allow_html=True)
    
    col_spacer, col_content, col_spacer2 = st.columns([0.5, 2, 0.5])
    
    with col_content:
        recent_items = [
            {"title": "Gerar um ebook sobre maternidade real", "meta": "Gerar • há 2 dias"},
            {"title": "Criar conteúdo sobre bem-estar mental", "meta": "Colar texto • há 5 dias"},
            {"title": "Transformar artigo em ebook profissional", "meta": "Importar • há 1 semana"},
        ]
        
        for item in recent_items:
            st.markdown(f"""
                <div style="
                    background: white;
                    border-radius: 12px;
                    padding: 20px 24px;
                    box-shadow: 0 5px 15px rgba(0, 0, 0, 0.08);
                    margin-bottom: 12px;
                    cursor: pointer;
                    transition: all 0.3s ease;
                ">
                    <p style="font-size: 16px; font-weight: 600; color: #0c2340; margin: 0 0 4px 0;">{item['title']}</p>
                    <p style="font-size: 12px; color: #94a3b8; margin: 0;">{item['meta']}</p>
                </div>
            """, unsafe_allow_html=True)

# PAGE: CRIAR EBOOK - COM NOVAS FUNCOES
elif st.session_state.page == "criar":
    if st.button("← Voltar", key="back_criar"):
        st.session_state.page = "onboarding"
        st.rerun()
    
    st.markdown("""
        <div class="header-premium">
            <h1>📚 Gerar Ebook com IA</h1>
            <p>O que você gostaria de criar hoje?</p>
        </div>
    """, unsafe_allow_html=True)
    
    st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
    
    # SELETOR DE TIPO
    st.markdown('<div class="section-header"><h2>1️⃣ Tipo de Documento</h2></div>', unsafe_allow_html=True)
    
    col1, col2, col3, col4, col5 = st.columns(5, gap="small")
    
    tipos_lista = list(TIPOS_DOCUMENTO.items())
    
    with col1:
        if st.button(f"{tipos_lista[0][1]}\n{tipos_lista[0][0]}", use_container_width=True, key="type_0"):
            st.session_state.selected_type = tipos_lista[0][0]
    
    with col2:
        if st.button(f"{tipos_lista[1][1]}\n{tipos_lista[1][0]}", use_container_width=True, key="type_1"):
            st.session_state.selected_type = tipos_lista[1][0]
    
    with col3:
        if st.button(f"{tipos_lista[2][1]}\n{tipos_lista[2][0]}", use_container_width=True, key="type_2"):
            st.session_state.selected_type = tipos_lista[2][0]
    
    with col4:
        if st.button(f"{tipos_lista[3][1]}\n{tipos_lista[3][0]}", use_container_width=True, key="type_3"):
            st.session_state.selected_type = tipos_lista[3][0]
    
    with col5:
        if st.button(f"{tipos_lista[4][1]}\n{tipos_lista[4][0]}", use_container_width=True, key="type_4"):
            st.session_state.selected_type = tipos_lista[4][0]
    
    st.markdown(f'<div class="success-box"><strong>✅ Selecionado: {st.session_state.selected_type}</strong></div>', unsafe_allow_html=True)
    
    st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
    
    # CONFIGURACOES
    st.markdown('<div class="section-header"><h2>2️⃣ Configurações</h2></div>', unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        num_itens = st.selectbox("Quantidade", ["5", "10", "15", "20", "25"], key="qtd_criar")
    
    with col2:
        estilo_select = st.selectbox("Estilo", list(ESTILOS.keys()), key="estilo_criar")
    
    with col3:
        idioma_option = st.selectbox("Idioma", list(IDIOMAS.keys()), key="idioma_criar")
    
    with col4:
        regiao = st.selectbox("Região", ["Portugal", "Brasil", "Angola", "Moçambique"], key="regiao_criar")
    
    st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
    
    # INPUT TEXT
    st.markdown('<div class="section-header"><h2>3️⃣ Descrição</h2></div>', unsafe_allow_html=True)
    prompt_user = st.text_area(
        "O que você quer criar?",
        placeholder="Ex: Brainstorming e criação de novas ideias para criação de conteúdo",
        height=120
    )
    
    st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
    
    # FORMATOS
    st.markdown('<div class="section-header"><h2>4️⃣ Formatos de Download</h2></div>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        formato_txt = st.checkbox("📄 TXT", value=True, key="txt_criar")
    with col2:
        formato_word = st.checkbox("📋 Word", value=True, key="word_criar")
    with col3:
        formato_pdf = st.checkbox("🎨 PDF", value=True, key="pdf_criar")
    
    st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
    
    # BOTAO GERAR
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        if st.button("✨ Gerar contorno", use_container_width=True, key="gen_criar"):
            if prompt_user.strip():
                try:
                    api_key = st.secrets["GEMINI_API_KEY"]
                    genai.configure(api_key=api_key)
                    models = [m.name for m in genai.list_models()]
                    model_name = [m for m in models if "gemini" in m.lower()][0]
                    model = genai.GenerativeModel(model_name)
                    
                    progress = st.progress(0)
                    status = st.empty()
                    
                    status.text("🔄 Gerando conteúdo...")
                    progress.progress(30)
                    
                    prompt_final = f"Crie um {st.session_state.selected_type} em {IDIOMAS[idioma_option]} com {num_itens} itens sobre: {prompt_user}\n\nEstilo: {ESTILOS[estilo_select]}\nRegião: {regiao}"
                    response = model.generate_content(prompt_final)
                    content = response.text
                    
                    progress.progress(70)
                    status.text("📸 Gerando sugestões...")
                    
                    response_img = model.generate_content(f"Crie 10 prompts de imagem para: {prompt_user}")
                    sugestoes = response_img.text
                    
                    progress.progress(100)
                    st.markdown('<div class="success-box"><strong>✅ Conteúdo gerado!</strong></div>', unsafe_allow_html=True)
                    
                    with st.expander("Ver conteúdo", expanded=False):
                        st.write(content)
                    
                    st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
                    fazer_download_ebook("Conteudo-Gerado", "Público geral", IDIOMAS[idioma_option], regiao, content, sugestoes, formato_txt, formato_word, formato_pdf)
                except Exception as e:
                    st.error(f"Erro: {str(e)}")
            else:
                st.warning("Descreva o que quer criar!")

# PAGE: COLAR TEXTO
elif st.session_state.page == "colar":
    if st.button("← Voltar", key="back_colar"):
        st.session_state.page = "onboarding"
        st.rerun()
    
    st.markdown("""
        <div class="header-premium">
            <h1>✏️ Transformar Texto em Ebook</h1>
            <p>Cole seu conteúdo e transforme em profissional</p>
        </div>
    """, unsafe_allow_html=True)
    
    st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
    
    conteudo_colado = st.text_area("Cole seu conteúdo", placeholder="Cole aqui anotações, artigos, ou qualquer texto", height=250)
    
    st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        titulo_ebook = st.text_input("Título do Ebook", "Meu Ebook Profissional")
    with col2:
        idioma_option = st.selectbox("Idioma", list(IDIOMAS.keys()), key="idioma_colar")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        formato_txt = st.checkbox("📄 TXT", value=True, key="txt_colar")
    with col2:
        formato_word = st.checkbox("📋 Word", value=True, key="word_colar")
    with col3:
        formato_pdf = st.checkbox("🎨 PDF", value=True, key="pdf_colar")
    
    st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
    
    if st.button("🚀 TRANSFORMAR", use_container_width=True, key="gen_colar"):
        if conteudo_colado.strip():
            try:
                api_key = st.secrets["GEMINI_API_KEY"]
                genai.configure(api_key=api_key)
                models = [m.name for m in genai.list_models()]
                model_name = [m for m in models if "gemini" in m.lower()][0]
                model = genai.GenerativeModel(model_name)
                
                progress = st.progress(0)
                status = st.empty()
                
                status.text("🔄 Estruturando conteúdo...")
                progress.progress(30)
                
                prompt = f"Transforme este texto em um ebook bem estruturado em {IDIOMAS[idioma_option]}, com capítulos, introdução e conclusão:\n\n{conteudo_colado}"
                response = model.generate_content(prompt)
                content = response.text
                
                progress.progress(70)
                status.text("📸 Gerando sugestões...")
                
                response_img = model.generate_content(f"Crie 10 prompts de imagem para este conteúdo: {conteudo_colado[:100]}")
                sugestoes = response_img.text
                
                progress.progress(100)
                st.markdown('<div class="success-box"><strong>✅ Transformado com sucesso!</strong></div>', unsafe_allow_html=True)
                
                with st.expander("Ver conteúdo", expanded=False):
                    st.write(content)
                
                st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
                fazer_download_ebook(titulo_ebook, "Leitores", IDIOMAS[idioma_option], "Portugal", content, sugestoes, formato_txt, formato_word, formato_pdf)
            except Exception as e:
                st.error(f"Erro: {str(e)}")
        else:
            st.warning("Cole seu conteúdo para começar!")

# PAGE: USAR MODELO
elif st.session_state.page == "modelo":
    if st.button("← Voltar", key="back_modelo"):
        st.session_state.page = "onboarding"
        st.rerun()
    
    st.markdown("""
        <div class="header-premium">
            <h1>📋 Usar Modelo</h1>
            <p>Escolha um modelo e customize com seus dados</p>
        </div>
    """, unsafe_allow_html=True)
    
    st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
    st.markdown('<div class="section-header"><h2>Escolha um Modelo</h2></div>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📖 Guia Completo", use_container_width=True):
            st.session_state.selected_template = "📖 Guia Completo"
    
    with col2:
        if st.button("📋 Check-list", use_container_width=True):
            st.session_state.selected_template = "📋 Check-list"
    
    with col3:
        if st.button("💡 Dicas Rápidas", use_container_width=True):
            st.session_state.selected_template = "💡 Dicas Rápidas"
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📊 Análise Dados", use_container_width=True):
            st.session_state.selected_template = "📊 Análise Dados"
    
    with col2:
        if st.button("🎯 Plano de Ação", use_container_width=True):
            st.session_state.selected_template = "🎯 Plano de Ação"
    
    with col3:
        if st.button("❓ FAQ", use_container_width=True):
            st.session_state.selected_template = "❓ FAQ"
    
    if hasattr(st.session_state, 'selected_template'):
        selected_template = st.session_state.selected_template
        st.markdown(f'<div class="success-box"><strong>✅ Modelo selecionado: {selected_template}</strong></div>', unsafe_allow_html=True)
        
        st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
        st.markdown('<div class="section-header"><h2>Customize seu Ebook</h2></div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        with col1:
            tema_option = st.selectbox("Tema", list(TEMAS.keys()), key="tema_modelo")
            tema_nome = tema_option.split(" ", 1)[1] if " " in tema_option else tema_option
        with col2:
            idioma_option = st.selectbox("Idioma", list(IDIOMAS.keys()), key="idioma_modelo")
        
        audience = st.text_input("Público-alvo", "Público geral")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            formato_txt = st.checkbox("📄 TXT", value=True, key="txt_modelo")
        with col2:
            formato_word = st.checkbox("📋 Word", value=True, key="word_modelo")
        with col3:
            formato_pdf = st.checkbox("🎨 PDF", value=True, key="pdf_modelo")
        
        st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
        
        if st.button("🚀 GERAR A PARTIR DO MODELO", use_container_width=True, key="gen_modelo"):
            try:
                api_key = st.secrets["GEMINI_API_KEY"]
                genai.configure(api_key=api_key)
                models = [m.name for m in genai.list_models()]
                model_name = [m for m in models if "gemini" in m.lower()][0]
                model = genai.GenerativeModel(model_name)
                
                progress = st.progress(0)
                status = st.empty()
                
                status.text("🔄 Gerando ebook a partir do modelo...")
                progress.progress(30)
                
                prompt = f"Crie um ebook no formato '{selected_template}' sobre '{TEMAS[tema_option]}' para '{audience}' em {IDIOMAS[idioma_option]}"
                response = model.generate_content(prompt)
                content = response.text
                
                progress.progress(70)
                status.text("📸 Gerando sugestões...")
                
                response_img = model.generate_content(f"Crie 10 prompts de imagem para: {TEMAS[tema_option]}")
                sugestoes = response_img.text
                
                progress.progress(100)
                st.markdown('<div class="success-box"><strong>✅ Ebook criado!</strong></div>', unsafe_allow_html=True)
                
                with st.expander("Ver conteúdo", expanded=False):
                    st.write(content)
                
                st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
                fazer_download_ebook(tema_nome, audience, IDIOMAS[idioma_option], "Portugal", content, sugestoes, formato_txt, formato_word, formato_pdf)
            except Exception as e:
                st.error(f"Erro: {str(e)}")

# PAGE: IMPORTAR
elif st.session_state.page == "importar":
    if st.button("← Voltar", key="back_importar"):
        st.session_state.page = "onboarding"
        st.rerun()
    
    st.markdown("""
        <div class="header-premium">
            <h1>📂 Importar & Transformar</h1>
            <p>Upload de arquivo ou URL</p>
        </div>
    """, unsafe_allow_html=True)
    
    st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
    st.markdown('<div class="section-header"><h2>Escolha uma opção</h2></div>', unsafe_allow_html=True)
    
    import_type = st.radio("Como quer importar?", ["📤 Upload de Arquivo", "🔗 URL"], horizontal=True)
    
    st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
    
    if import_type == "📤 Upload de Arquivo":
        arquivo = st.file_uploader("Selecione um arquivo (TXT, DOCX, PDF)", type=["txt", "docx", "pdf"])
        
        if arquivo:
            st.success(f"✅ Arquivo carregado: {arquivo.name}")
            
            try:
                if arquivo.name.endswith('.txt'):
                    conteudo_arquivo = arquivo.getvalue().decode()
                elif arquivo.name.endswith('.docx'):
                    from docx import Document as DocRead
                    doc = DocRead(arquivo)
                    conteudo_arquivo = "\n".join([p.text for p in doc.paragraphs])
                else:
                    conteudo_arquivo = "[Arquivo PDF - extrair manualmente]"
            except:
                conteudo_arquivo = ""
            
            idioma_option = st.selectbox("Idioma", list(IDIOMAS.keys()), key="idioma_import")
            titulo = st.text_input("Título do Ebook", arquivo.name.split('.')[0])
            
            col1, col2, col3 = st.columns(3)
            with col1:
                formato_txt = st.checkbox("📄 TXT", value=True, key="txt_import")
            with col2:
                formato_word = st.checkbox("📋 Word", value=True, key="word_import")
            with col3:
                formato_pdf = st.checkbox("🎨 PDF", value=True, key="pdf_import")
            
            st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
            
            if st.button("🚀 TRANSFORMAR ARQUIVO", use_container_width=True, key="gen_import_file"):
                try:
                    api_key = st.secrets["GEMINI_API_KEY"]
                    genai.configure(api_key=api_key)
                    models = [m.name for m in genai.list_models()]
                    model_name = [m for m in models if "gemini" in m.lower()][0]
                    model = genai.GenerativeModel(model_name)
                    
                    progress = st.progress(0)
                    status = st.empty()
                    
                    status.text("🔄 Processando arquivo...")
                    progress.progress(30)
                    
                    prompt = f"Transforme este conteúdo em um ebook profissional bem estruturado em {IDIOMAS[idioma_option]}:\n\n{conteudo_arquivo[:2000]}"
                    response = model.generate_content(prompt)
                    content = response.text
                    
                    progress.progress(70)
                    status.text("📸 Gerando sugestões...")
                    
                    response_img = model.generate_content(f"Crie 10 prompts de imagem para este ebook")
                    sugestoes = response_img.text
                    
                    progress.progress(100)
                    st.markdown('<div class="success-box"><strong>✅ Arquivo transformado!</strong></div>', unsafe_allow_html=True)
                    
                    with st.expander("Ver conteúdo", expanded=False):
                        st.write(content)
                    
                    st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
                    fazer_download_ebook(titulo, "Leitores", IDIOMAS[idioma_option], "Portugal", content, sugestoes, formato_txt, formato_word, formato_pdf)
                except Exception as e:
                    st.error(f"Erro: {str(e)}")
    else:
        url = st.text_input("Cole a URL")
        
        if url:
            idioma_option = st.selectbox("Idioma", list(IDIOMAS.keys()), key="idioma_import_url")
            titulo = st.text_input("Título do Ebook", "Ebook Importado")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                formato_txt = st.checkbox("📄 TXT", value=True, key="txt_import_url")
            with col2:
                formato_word = st.checkbox("📋 Word", value=True, key="word_import_url")
            with col3:
                formato_pdf = st.checkbox("🎨 PDF", value=True, key="pdf_import_url")
            
            st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
            
            if st.button("🚀 TRANSFORMAR URL", use_container_width=True, key="gen_import_url"):
                try:
                    api_key = st.secrets["GEMINI_API_KEY"]
                    genai.configure(api_key=api_key)
                    models = [m.name for m in genai.list_models()]
                    model_name = [m for m in models if "gemini" in m.lower()][0]
                    model = genai.GenerativeModel(model_name)
                    
                    progress = st.progress(0)
                    status = st.empty()
                    
                    status.text("🔄 Acessando URL...")
                    progress.progress(20)
                    
                    try:
                        response = requests.get(url, timeout=10)
                        conteudo_url = response.text[:3000]
                    except:
                        conteudo_url = f"[URL: {url}]"
                    
                    progress.progress(40)
                    status.text("🔄 Gerando ebook...")
                    
                    prompt = f"Transforme o conteúdo desta URL em um ebook profissional em {IDIOMAS[idioma_option]}: {conteudo_url}"
                    response = model.generate_content(prompt)
                    content = response.text
                    
                    progress.progress(70)
                    status.text("📸 Gerando sugestões...")
                    
                    response_img = model.generate_content("Crie 10 prompts de imagem para este ebook")
                    sugestoes = response_img.text
                    
                    progress.progress(100)
                    st.markdown('<div class="success-box"><strong>✅ URL transformada!</strong></div>', unsafe_allow_html=True)
                    
                    with st.expander("Ver conteúdo", expanded=False):
                        st.write(content)
                    
                    st.markdown('<div class="divider-premium"></div>', unsafe_allow_html=True)
                    fazer_download_ebook(titulo, "Leitores", IDIOMAS[idioma_option], "Portugal", content, sugestoes, formato_txt, formato_word, formato_pdf)
                except Exception as e:
                    st.error(f"Erro: {str(e)}")

# FOOTER
st.markdown("""
    <div class="footer-custom">
        <p><strong>Luciana Britto | L&B Marketing — Estratégias de Valor</strong></p>
        <p>© 2026 • Ferramenta Premium de IA para Empreendoras</p>
        <p style="font-size: 12px; opacity: 0.7;">Powered by Google Gemini AI</p>
    </div>
""", unsafe_allow_html=True)
